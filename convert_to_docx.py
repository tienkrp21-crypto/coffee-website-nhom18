from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create Word document
doc = Document()

# Add title
title = doc.add_heading('Business Requirements Document (BRD)', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle = doc.add_heading('Enterprise Asset Management System (KLM-EAM)', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add metadata
meta = doc.add_paragraph()
meta.add_run('Ngày: ').bold = True
meta.add_run('04/05/2026 | ')
meta.add_run('Version: ').bold = True
meta.add_run('1.0 | ')
meta.add_run('Người chuẩn bị: ').bold = True
meta.add_run('Trương Hùng Dũng')

doc.add_paragraph()

# Content sections
doc.add_heading('1. TỔNG QUAN DỰ ÁN', level=1)

doc.add_heading('1.1 Mô tả tổng quát', level=2)
doc.add_paragraph('Enterprise Asset Management System (KLM-EAM) là giải pháp số hóa toàn diện cho việc quản lý tài sản và thiết bị đo lường nội bộ của Công ty Kim Long. Hệ thống cung cấp nền tảng thống nhất để:')
doc.add_paragraph('Theo dõi vòng đời hoàn chỉnh của thiết bị', style='List Bullet')
doc.add_paragraph('Tự động hóa quy trình hiệu chuẩn', style='List Bullet')
doc.add_paragraph('Kiểm soát chặt chẽ quá trình luân chuyển thiết bị giữa Lab và hiện trường', style='List Bullet')
doc.add_paragraph('Cung cấp dashboard tập trung theo dõi tình trạng tài sản real-time', style='List Bullet')

doc.add_heading('1.2 Bối cảnh kinh doanh', level=2)
doc.add_paragraph('Tình trạng hiện tại: Công ty Kim Long hiện sử dụng quy trình thủ công để quản lý tài sản')
doc.add_paragraph('Thách thức: Khó quản lý hàng trăm thiết bị tại nhiều chi nhánh (HCM, Cần Thơ), dễ thất lạc dữ liệu')
doc.add_paragraph('Mục tiêu: Đạt độ minh bạch 100% trong quá trình mượn/trả, giảm 80% lỗi thủ công')

# Stakeholders table
doc.add_heading('1.3 Các bên liên quan (Stakeholders)', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Vai trò'
hdr_cells[1].text = 'Trách nhiệm'
hdr_cells[2].text = 'Quyền hạn'

table.rows[1].cells[0].text = 'Admin'
table.rows[1].cells[1].text = 'Quản lý người dùng, cấu hình hệ thống'
table.rows[1].cells[2].text = 'Toàn quyền'

table.rows[2].cells[0].text = 'QC'
table.rows[2].cells[1].text = 'Xác nhận xuất/nhập thiết bị'
table.rows[2].cells[2].text = 'Xác nhận, kiểm tra'

table.rows[3].cells[0].text = 'KT'
table.rows[3].cells[1].text = 'Mượn/trả thiết bị'
table.rows[3].cells[2].text = 'Mượn, trả, xem lịch sử'

table.rows[4].cells[0].text = 'Sales'
table.rows[4].cells[1].text = 'Xem thông tin tài sản (chỉ đọc)'
table.rows[4].cells[2].text = 'Xem báo cáo'

table.rows[5].cells[0].text = 'Quản lý Kho'
table.rows[5].cells[1].text = 'Quản lý vị trí, cấu hình trạng thái'
table.rows[5].cells[2].text = 'Quản lý master data'

doc.add_page_break()

doc.add_heading('2. PHẠM VI DỰ ÁN', level=1)

doc.add_heading('2.1 Chức năng trong phạm vi (In Scope)', level=2)
doc.add_heading('A. Quản lý Master Data', level=3)
doc.add_paragraph('Quản lý danh mục thiết bị', style='List Bullet')
doc.add_paragraph('Tạo/Cập nhật/Xóa thiết bị', style='List Bullet 2')
doc.add_paragraph('Quản lý chu kỳ hiệu chuẩn', style='List Bullet 2')

doc.add_heading('B. Quản lý Hiệu chuẩn (Calibration)', level=3)
doc.add_paragraph('Ghi nhận ngày hiệu chuẩn gần nhất', style='List Bullet')
doc.add_paragraph('Tính toán tự động ngày hiệu chuẩn tiếp theo', style='List Bullet')
doc.add_paragraph('Cảnh báo: Màu đỏ nếu quá hạn, màu cam nếu còn <= 30 ngày', style='List Bullet')

doc.add_heading('C. Quy trình Luân chuyển Thiết bị (Transfer Process)', level=3)
doc.add_paragraph('KT mượn thiết bị (quét QR, chọn lý do, chọn ngày trả)', style='List Number')
doc.add_paragraph('QC xác nhận xuất (kiểm tra, xác nhận hoặc từ chối)', style='List Number')
doc.add_paragraph('KT trả thiết bị (quét QR, ghi tình trạng)', style='List Number')
doc.add_paragraph('QC xác nhận nhập (kiểm tra, cập nhật trạng thái)', style='List Number')

doc.add_page_break()

doc.add_heading('4. YÊU CẦU CHỨC NĂNG (FUNCTIONAL REQUIREMENTS)', level=1)

fr_data = [
    ['FR-AUTH-001', 'MUST', 'Xác thực', 'Đăng nhập bằng email/username', 'Tất cả'],
    ['FR-AUTH-002', 'MUST', 'Xác thực', 'Quên mật khẩu qua email', 'KT, QC, Admin'],
    ['FR-EQUIP-001', 'MUST', 'Quản lý Thiết bị', 'Admin có thể tạo/sửa/xóa thiết bị', 'Admin'],
    ['FR-EQUIP-002', 'MUST', 'Quản lý Thiết bị', 'Mỗi thiết bị có: Asset Code, Model, Serial, Manufacturer, Location, Status', 'Admin, QC'],
    ['FR-CAL-001', 'MUST', 'Hiệu chuẩn', 'Tự động tính ngày hiệu chuẩn tiếp theo', 'Admin, QC'],
    ['FR-CAL-002', 'MUST', 'Hiệu chuẩn', 'Cảnh báo thiết bị sắp đến hạn (<= 30 ngày)', 'Admin, QC'],
    ['FR-TRF-001', 'MUST', 'Luân chuyển', 'KT quét mã QR để mượn thiết bị', 'KT'],
    ['FR-TRF-003', 'MUST', 'Luân chuyển', 'QC xác nhận xuất: kiểm tra, xác nhận/từ chối', 'QC'],
    ['FR-TRF-005', 'MUST', 'Luân chuyển', 'KT trả thiết bị, ghi tình trạng', 'KT'],
    ['FR-TRF-006', 'MUST', 'Luân chuyển', 'QC xác nhận nhập: kiểm tra, cập nhật status', 'QC'],
    ['FR-DASH-001', 'MUST', 'Dashboard', 'Hiển thị KPI: Tổng thiết bị, ở Lab, On-site, Bảo trì', 'Admin, Quản lý'],
    ['FR-DASH-002', 'MUST', 'Dashboard', 'Cảnh báo thiết bị sắp đến hạn', 'Admin, Quản lý'],
    ['FR-DASH-003', 'MUST', 'Dashboard', 'Bảng danh sách thiết bị với cột: Asset, Model, Location, Status', 'Admin, Quản lý'],
    ['FR-DASH-004', 'MUST', 'Dashboard', 'Filter bảng danh sách theo Location và Status', 'Admin, Quản lý'],
]

table = doc.add_table(rows=len(fr_data)+1, cols=5)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Ưu tiên'
hdr_cells[2].text = 'Module'
hdr_cells[3].text = 'Mô tả'
hdr_cells[4].text = 'Stakeholders'

for idx, row_data in enumerate(fr_data, 1):
    for col, cell_data in enumerate(row_data):
        table.rows[idx].cells[col].text = str(cell_data)

doc.add_page_break()

doc.add_heading('5. YÊU CẦU PHI CHỨC NĂNG (NON-FUNCTIONAL REQUIREMENTS)', level=1)

nfr_data = [
    ['NFR-PERF-001', 'Performance', 'Thời gian tải trang < 3 giây', 'Page load < 3s'],
    ['NFR-PERF-002', 'Performance', 'Hỗ trợ tối thiểu 100 concurrent users', 'Load test pass'],
    ['NFR-SEC-001', 'Security', 'Mã hóa mật khẩu (bcrypt, salt 10+)', 'Passwords hashed'],
    ['NFR-SEC-002', 'Security', 'Token-based authentication (JWT)', 'JWT implemented'],
    ['NFR-SEC-003', 'Security', 'HTTPS/TLS bắt buộc', 'SSL certificate'],
    ['NFR-SCAL-001', 'Scalability', 'Database indexing', 'Indexes created'],
    ['NFR-AUDIT-001', 'Audit', 'Ghi nhận mọi thay đổi dữ liệu', 'Audit table created'],
]

table = doc.add_table(rows=len(nfr_data)+1, cols=4)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Loại'
hdr_cells[2].text = 'Mô tả'
hdr_cells[3].text = 'Tiêu chí chấp nhận'

for idx, row_data in enumerate(nfr_data, 1):
    for col, cell_data in enumerate(row_data):
        table.rows[idx].cells[col].text = str(cell_data)

doc.add_page_break()

doc.add_heading('9. TIMELINE & PHÂN GIA ĐOẠN', level=1)

timeline_data = [
    ['Phase', 'Tháng', 'Milestone', 'Deliverable'],
    ['Phase 1', 'T5-T6/2026', 'Requirements, DB Design', 'BRD, Schema, API Spec'],
    ['Phase 2', 'T7-T8/2026', 'Development', 'Backend APIs, Frontend UI'],
    ['Phase 3', 'T8-T9/2026', 'Testing & UAT', 'Test Reports, UAT Results'],
    ['Phase 4', 'T9/2026', 'Deployment', 'Production Release'],
]

table = doc.add_table(rows=len(timeline_data), cols=4)
table.style = 'Light Grid Accent 1'

for idx, row_data in enumerate(timeline_data):
    for col, cell_data in enumerate(row_data):
        table.rows[idx].cells[col].text = cell_data

doc.add_paragraph()
doc.add_paragraph('END OF DOCUMENT')

# Save file
output_path = r'E:\BRD_KLM_EAM_Ver1.0.docx'
doc.save(output_path)
print(f"File saved successfully to: {output_path}")
